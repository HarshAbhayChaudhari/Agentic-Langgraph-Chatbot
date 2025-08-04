import os
import re
from typing import List, Dict, Optional
from datetime import datetime
import tempfile

# PDF parsing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class FileParser:
    """
    Universal file parser that can handle TXT, PDF, and DOCX files
    """
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._parse_txt,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx
        }
    
    def parse_file(self, file_path: str) -> List[Dict]:
        """
        Parse file based on its extension
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            List of parsed messages/documents
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        parser_method = self.supported_formats[file_extension]
        return parser_method(file_path)
    
    def _parse_txt(self, file_path: str) -> List[Dict]:
        """
        Parse TXT file (WhatsApp chat export format)
        """
        messages = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
            # WhatsApp chat export pattern
            message_pattern = r'\[(\d{1,2}/\d{1,2}/\d{4}),\s*(\d{1,2}:\d{2}:\d{2})\]\s*(.+?):\s*(.+)'
            matches = re.findall(message_pattern, content, re.MULTILINE)
            
            for match in matches:
                date_str, time_str, sender, message = match
                
                try:
                    datetime_obj = datetime.strptime(f"{date_str} {time_str}", "%d/%m/%Y %H:%M:%S")
                except ValueError:
                    continue
                
                cleaned_message = self._clean_message(message)
                
                if cleaned_message.strip():
                    messages.append({
                        'date': datetime_obj,
                        'sender': sender.strip(),
                        'message': cleaned_message,
                        'timestamp': datetime_obj.isoformat(),
                        'source': 'whatsapp_chat'
                    })
                    
        except Exception as e:
            print(f"Error parsing TXT file: {str(e)}")
            # If WhatsApp format fails, treat as plain text
            return self._parse_plain_text(file_path)
            
        return messages
    
    def _parse_plain_text(self, file_path: str) -> List[Dict]:
        """
        Parse plain text file (fallback for non-WhatsApp TXT files)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Split content into paragraphs
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            messages = []
            for i, paragraph in enumerate(paragraphs):
                if len(paragraph) > 10:  # Only include substantial paragraphs
                    messages.append({
                        'date': datetime.now(),
                        'sender': 'Document',
                        'message': paragraph,
                        'timestamp': datetime.now().isoformat(),
                        'source': 'plain_text',
                        'paragraph_id': i
                    })
            
            return messages
            
        except Exception as e:
            print(f"Error parsing plain text: {str(e)}")
            return []
    
    def _parse_pdf(self, file_path: str) -> List[Dict]:
        """
        Parse PDF file
        """
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing libraries not available. Install PyPDF2 and pdfplumber.")
        
        messages = []
        
        try:
            # Try pdfplumber first (better text extraction)
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        # Split text into paragraphs
                        paragraphs = [p.strip() for p in text.split('\n') if p.strip() and len(p.strip()) > 10]
                        
                        for i, paragraph in enumerate(paragraphs):
                            messages.append({
                                'date': datetime.now(),
                                'sender': f'PDF_Page_{page_num + 1}',
                                'message': paragraph,
                                'timestamp': datetime.now().isoformat(),
                                'source': 'pdf',
                                'page': page_num + 1,
                                'paragraph_id': i
                            })
            
            # If pdfplumber didn't extract much, try PyPDF2
            if len(messages) < 2:
                messages = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            paragraphs = [p.strip() for p in text.split('\n') if p.strip() and len(p.strip()) > 10]
                            
                            for i, paragraph in enumerate(paragraphs):
                                messages.append({
                                    'date': datetime.now(),
                                    'sender': f'PDF_Page_{page_num + 1}',
                                    'message': paragraph,
                                    'timestamp': datetime.now().isoformat(),
                                    'source': 'pdf',
                                    'page': page_num + 1,
                                    'paragraph_id': i
                                })
                                
        except Exception as e:
            print(f"Error parsing PDF file: {str(e)}")
            return []
        
        return messages
    
    def _parse_docx(self, file_path: str) -> List[Dict]:
        """
        Parse DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing library not available. Install python-docx.")
        
        messages = []
        
        try:
            doc = Document(file_path)
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text and len(text) > 10:  # Only include substantial paragraphs
                    messages.append({
                        'date': datetime.now(),
                        'sender': 'Document',
                        'message': text,
                        'timestamp': datetime.now().isoformat(),
                        'source': 'docx',
                        'paragraph_id': para_num
                    })
            
            # Also extract text from tables
            for table_num, table in enumerate(doc.tables):
                for row_num, row in enumerate(table.rows):
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text and len(row_text) > 10:
                        messages.append({
                            'date': datetime.now(),
                            'sender': f'Table_{table_num + 1}',
                            'message': row_text,
                            'timestamp': datetime.now().isoformat(),
                            'source': 'docx_table',
                            'table_id': table_num + 1,
                            'row_id': row_num + 1
                        })
                        
        except Exception as e:
            print(f"Error parsing DOCX file: {str(e)}")
            return []
        
        return messages
    
    def _clean_message(self, message: str) -> str:
        """
        Clean and preprocess message text
        """
        # Remove common WhatsApp artifacts
        cleaned = message.strip()
        
        # Remove media messages
        if any(media_indicator in cleaned.lower() for media_indicator in [
            '<media omitted>', 'image omitted', 'video omitted', 
            'audio omitted', 'document omitted', 'location omitted'
        ]):
            return ""
        
        # Remove excessive whitespace
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        return cleaned.strip()
    
    def get_file_info(self, file_path: str) -> Dict:
        """
        Get information about the file being parsed
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        file_size = os.path.getsize(file_path)
        
        return {
            'filename': os.path.basename(file_path),
            'extension': file_extension,
            'size_bytes': file_size,
            'size_mb': round(file_size / (1024 * 1024), 2),
            'supported': file_extension in self.supported_formats
        } 